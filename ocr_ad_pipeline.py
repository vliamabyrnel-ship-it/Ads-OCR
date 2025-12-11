#!/usr/bin/env python3

from __future__ import annotations

import argparse
import json
import logging
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

import cv2
import numpy as np
from PIL import Image

# --- Optional / heavy dependencies ---------------------------------------
# These imports are guarded so the script fails gracefully and logs warnings
# instead of just crashing.

try:
    import easyocr  # Deep-learning OCR engine (CRAFT + CRNN)
except ImportError:
    easyocr = None

try:
    import pytesseract  # Wrapper around system Tesseract binary
except ImportError:
    pytesseract = None

try:
    from transformers import pipeline as hf_pipeline  # HuggingFace summarizer
except ImportError:
    hf_pipeline = None

try:
    from docx import Document  # python-docx for Word creation
    from docx.shared import Inches
except ImportError:
    Document = None
    Inches = None


# -------------------------------------------------------------------------
#                               DATA CLASSES
# -------------------------------------------------------------------------

@dataclass
class OCRConfig:
    """
    Configuration for the OCR pipeline.

    Having this in one place makes it easy to tweak things, or even
    expose them as command line flags in the future.
    """
    languages: Tuple[str, ...] = ("en",)
    resize_factor: float = 2.0
    min_confidence: float = 0.4
    use_tesseract_backup: bool = True
    use_ai_summarizer: bool = True
    summarizer_model: str = "facebook/bart-large-cnn"
    gpu: bool = True  # we'll still safely fall back to CPU if needed


@dataclass
class TextBox:
    """
    Container for one recognized text fragment.

    Attributes
    ----------
    text : str
        The cleaned text string.
    confidence : float
        Confidence on [0, 1].
    bbox : List[Tuple[int, int]]
        List of 4 points defining the bounding box (x, y).
    source : str
        A tag indicating which engine / pre-processing variant produced this box.
    """
    text: str
    confidence: float
    bbox: List[Tuple[int, int]]
    source: str


# -------------------------------------------------------------------------
#                                UTILITIES
# -------------------------------------------------------------------------

def setup_logging(verbose: bool = False) -> None:
    """
    Configure logging for the script. Call this early.

    Parameters
    ----------
    verbose : bool
        If True, log at DEBUG level; otherwise INFO.
    """
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format="%(asctime)s [%(levelname)s] %(message)s",
    )


def normalize_text(text: str) -> str:
    """
    Normalize raw OCR output a bit:
        - strip leading/trailing whitespace
        - collapse internal whitespace to single spaces
    """
    if not text:
        return ""
    cleaned = " ".join(text.split())
    return cleaned.strip()


def load_image(path: Path, cfg: OCRConfig) -> np.ndarray:
    """
    Load an image from disk and optionally upscale it.

    We use OpenCV here because most pre-processing later is also OpenCV-based.
    """
    img = cv2.imread(str(path))
    if img is None:
        raise ValueError(f"Could not read image from {path}")

    if cfg.resize_factor != 1.0:
        img = cv2.resize(
            img,
            dsize=None,
            fx=cfg.resize_factor,
            fy=cfg.resize_factor,
            interpolation=cv2.INTER_CUBIC,
        )

    return img


def generate_variants(bgr_img: np.ndarray) -> Dict[str, np.ndarray]:
    """
    Produce a small family of pre-processed image variants.

    Different types of text tend to pop out under different transformations,
    so we feed multiple variants into the OCR engines.

    Returns a dict mapping variant_name -> single-channel ndarray.
    """
    variants: Dict[str, np.ndarray] = {}

    gray = cv2.cvtColor(bgr_img, cv2.COLOR_BGR2GRAY)
    variants["gray"] = gray

    # Contrast Limited Adaptive Histogram Equalization
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    gray_clahe = clahe.apply(gray)
    variants["clahe"] = gray_clahe

    # Otsu binary thresholding
    _, thresh = cv2.threshold(
        gray_clahe, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
    )
    variants["binary"] = thresh
    variants["binary_inv"] = cv2.bitwise_not(thresh)

    # Blur + sharpen to reduce noise then enhance edges
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)
    sharpen_kernel = np.array(
        [[0, -1, 0],
         [-1, 5, -1],
         [0, -1, 0]]
    )
    sharpened = cv2.filter2D(blurred, -1, sharpen_kernel)
    variants["sharpened"] = sharpened

    return variants


# -------------------------------------------------------------------------
#                               OCR ENGINES
# -------------------------------------------------------------------------

def run_easyocr(
    img: np.ndarray,
    variant_name: str,
    cfg: OCRConfig,
    reader: "easyocr.Reader",
) -> List[TextBox]:
    """
    Run EasyOCR on the given image variant.

    Parameters
    ----------
    img : np.ndarray
        Grayscale or RGB image.
    variant_name : str
        Tag for debugging (e.g., "binary", "clahe").
    reader : easyocr.Reader
        Initialized EasyOCR reader.

    Returns
    -------
    List[TextBox]
    """
    text_boxes: List[TextBox] = []

    # EasyOCR accepts grayscale, but supports color as well.
    # We'll just feed the variant as-is.
    try:
        output = reader.readtext(img, detail=1, paragraph=False)
    except Exception as exc:
        logging.warning(f"EasyOCR failed on variant '{variant_name}': {exc}")
        return text_boxes

    for bbox, text, conf in output:
        cleaned = normalize_text(text)
        if not cleaned:
            continue
        if conf < cfg.min_confidence:
            continue

        tb = TextBox(
            text=cleaned,
            confidence=float(conf),
            bbox=[(int(x), int(y)) for x, y in bbox],
            source=f"easyocr:{variant_name}",
        )
        text_boxes.append(tb)

    return text_boxes


def run_tesseract(
    img: np.ndarray,
    variant_name: str,
    cfg: OCRConfig,
) -> List[TextBox]:
    """
    Run Tesseract OCR (via pytesseract) on the given image variant.

    Tesseract prefers high-contrast, single-channel images, which is exactly
    what our variants are.

    Returns a list of TextBox; bounding boxes are approximated based on
    Tesseract's word-level TSV.
    """
    results: List[TextBox] = []

    if pytesseract is None:
        # Library not available; nothing to do.
        return results

    pil_img = Image.fromarray(img)

    try:
        tsv = pytesseract.image_to_data(
            pil_img,
            output_type=pytesseract.Output.DICT,
            config="--oem 3 --psm 6",
        )
    except Exception as exc:
        logging.warning(f"Tesseract failed on variant '{variant_name}': {exc}")
        return results

    n_boxes = len(tsv["text"])
    for i in range(n_boxes):
        raw_text = tsv["text"][i]
        cleaned = normalize_text(raw_text)
        if not cleaned:
            continue

        conf = float(tsv["conf"][i])
        # Tesseract reports -1 for "no confidence"
        if conf < cfg.min_confidence * 100:
            continue

        x, y, w, h = (
            int(tsv["left"][i]),
            int(tsv["top"][i]),
            int(tsv["width"][i]),
            int(tsv["height"][i]),
        )
        bbox = [(x, y), (x + w, y), (x + w, y + h), (x, y + h)]

        results.append(
            TextBox(
                text=cleaned,
                confidence=conf / 100.0,
                bbox=bbox,
                source=f"tesseract:{variant_name}",
            )
        )

    return results


# -------------------------------------------------------------------------
#                         MERGING & TEXT ASSEMBLY
# -------------------------------------------------------------------------

def merge_text_boxes(boxes: Iterable[TextBox]) -> List[TextBox]:
    """
    Merge boxes from multiple engines/variants.

    Approach:
    ---------
    * Group by normalized (lowercased) text string.
    * For each group, keep the TextBox with the highest confidence.
    * Finally, sort by the vertical position (top y) to approximate reading order.
    """
    best_by_text: Dict[str, TextBox] = {}

    for tb in boxes:
        key = tb.text.lower()
        if key not in best_by_text or tb.confidence > best_by_text[key].confidence:
            best_by_text[key] = tb

    merged = list(best_by_text.values())
    merged.sort(key=lambda t: min(pt[1] for pt in t.bbox))

    return merged


def build_raw_text(boxes: List[TextBox]) -> str:
    """
    Build a single multi-line string from ordered TextBox entries.
    """
    return "\n".join(tb.text for tb in boxes)


# -------------------------------------------------------------------------
#                             SUMMARIZATION / AI
# -------------------------------------------------------------------------

def create_summarizer(cfg: OCRConfig):
    """
    Try to construct a HuggingFace summarization pipeline.

    Returns the pipeline object, or None if something goes wrong.
    """
    if not cfg.use_ai_summarizer:
        logging.info("AI summarizer disabled in config.")
        return None

    if hf_pipeline is None:
        logging.warning("transformers library not installed; "
                        "AI summarizer not available.")
        return None

    try:
        logging.info(f"Loading summarization model: {cfg.summarizer_model}")
        summarizer = hf_pipeline("summarization", model=cfg.summarizer_model)
        return summarizer
    except Exception as exc:
        logging.warning(f"Failed to initialize summarizer: {exc}")
        return None


def summarize_text(text: str, summarizer, min_length: int = 20, max_length: int = 120) -> str:
    """
    Summarize raw OCR text using the provided summarizer, or a heuristic fallback.
    """
    text = text.strip()
    if not text:
        return ""

    if summarizer is not None:
        try:
            # Note: HuggingFace uses token lengths, but character count roughly correlates.
            output = summarizer(
                text,
                max_length=max_length,
                min_length=min_length,
                do_sample=False,
            )
            summary = output[0]["summary_text"]
            return summary.strip()
        except Exception as exc:
            logging.warning(f"Summarization failed, falling back to heuristic: {exc}")

    # Heuristic fallback: choose lines with more uppercase letters (often ad headlines).
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return ""

    def score(line: str) -> float:
        if not line:
            return 0.0
        num_upper = sum(c.isupper() for c in line)
        return num_upper / len(line)

    sorted_lines = sorted(lines, key=score, reverse=True)
    top = sorted_lines[:10]  # arbitrary limit
    selected = [ln for ln in lines if ln in top]

    return "\n".join(selected)


# -------------------------------------------------------------------------
#                          WORD DOCUMENT GENERATION
# -------------------------------------------------------------------------

def create_word_report(
    results: List[Dict],
    output_path: Path,
) -> None:
    """
    Create a Word document with each image and its OCR results.

    Each image gets:
        - A heading with the file name
        - The image itself (resized to a reasonable width)
        - A "Raw OCR Text" section
        - A "Summarized Advertisement Text" section
    """
    if Document is None or Inches is None:
        raise RuntimeError(
            "python-docx is not installed, cannot create Word document."
        )

    doc = Document()

    for idx, res in enumerate(results, start=1):
        image_path: Path = res["image_path"]
        raw_text: str = res["raw_text"]
        summary: str = res["summarized_text"]

        doc.add_heading(f"Image {idx}: {image_path.name}", level=2)

        # Insert the image if it still exists (we check just in case).
        if image_path.exists():
            try:
                doc.add_picture(str(image_path), width=Inches(5.5))
            except Exception as exc:
                logging.warning(f"Failed to embed image {image_path}: {exc}")
        else:
            doc.add_paragraph(f"(Image file not found: {image_path})")

        doc.add_heading("Raw OCR Text", level=3)
        if raw_text.strip():
            doc.add_paragraph(raw_text)
        else:
            doc.add_paragraph("(No text detected)")

        doc.add_heading("Summarized Advertisement Text", level=3)
        if summary.strip():
            doc.add_paragraph(summary)
        else:
            doc.add_paragraph("(No summary available)")

        # Page break after each image except possibly the last.
        if idx != len(results):
            doc.add_page_break()

    doc.save(str(output_path))
    logging.info(f"Saved Word report to: {output_path}")


# -------------------------------------------------------------------------
#                           MAIN IMAGE PROCESSING
# -------------------------------------------------------------------------

def process_image(
    img_path: Path,
    cfg: OCRConfig,
    reader: Optional["easyocr.Reader"],
    summarizer,
) -> Dict:
    """
    Run the full OCR + summarization pipeline on a single image.

    Returns a dict suitable for JSON serialization and for use in the
    Word report generator.
    """
    logging.info(f"Processing image: {img_path}")

    bgr = load_image(img_path, cfg)
    variants = generate_variants(bgr)

    text_boxes: List[TextBox] = []

    # EasyOCR pass
    if reader is not None:
        for name, variant_img in variants.items():
            tbs = run_easyocr(variant_img, name, cfg, reader)
            text_boxes.extend(tbs)
    else:
        logging.warning("EasyOCR is not available; skipping EasyOCR pass.")

    # Tesseract backup
    if cfg.use_tesseract_backup and pytesseract is not None:
        for name, variant_img in variants.items():
            tbs = run_tesseract(variant_img, name, cfg)
            text_boxes.extend(tbs)
    elif cfg.use_tesseract_backup and pytesseract is None:
        logging.warning("pytesseract is not installed; Tesseract backup disabled.")

    merged_boxes = merge_text_boxes(text_boxes)
    raw_text = build_raw_text(merged_boxes)
    summarized_text = summarize_text(raw_text, summarizer=summarizer)

    # Prepare JSON-friendly structure
    json_ready = {
        "image_file": img_path.name,
        "config": asdict(cfg),
        "num_raw_boxes": len(text_boxes),
        "num_merged_boxes": len(merged_boxes),
        "raw_text": raw_text,
        "summarized_text": summarized_text,
        "boxes": [
            {
                "text": tb.text,
                "confidence": tb.confidence,
                "bbox": tb.bbox,
                "source": tb.source,
            }
            for tb in merged_boxes
        ],
    }
    return {
        "json": json_ready,
        "image_path": img_path,
        "raw_text": raw_text,
        "summarized_text": summarized_text,
    }


def find_images(image_dir: Path) -> List[Path]:
    """
    Return a sorted list of image files in the directory.

    Non-recursive on purpose: keeps behavior simple and predictable.
    """
    exts = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"}
    files = [p for p in image_dir.iterdir() if p.suffix.lower() in exts]
    files.sort()
    return files


# -------------------------------------------------------------------------
#                                 CLI / MAIN
# -------------------------------------------------------------------------

def parse_args() -> argparse.Namespace:
    """
    Parse command-line arguments for the script.

    We allow the caller to provide the image directory as a flag, but we
    will also prompt for an output directory at runtime as requested.
    """
    parser = argparse.ArgumentParser(
        description="OCR pipeline for advertisement images "
                    "(JSON + Word document output)."
    )
    parser.add_argument(
        "--image-dir",
        type=Path,
        default=Path(r"C:\Users\saddu\Desktop\VScode\Jassi\images"),
        help="Directory containing images to process.",
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Enable debug logging for troubleshooting.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    setup_logging(verbose=args.verbose)

    image_dir: Path = args.image_dir
    if not image_dir.exists() or not image_dir.is_dir():
        logging.error(f"Image directory does not exist or is not a directory: {image_dir}")
        return

    # Prompt the user for the output directory as requested.
    print("\nPlease enter the output folder where JSON reports and "
          "the Word document should be stored.")
    print("Example: C:\\Users\\saddu\\Desktop\\VScode\\Jassi\\output")
    out_dir_str = input("Output folder path: ").strip('"').strip()

    output_dir = Path(out_dir_str)
    output_dir.mkdir(parents=True, exist_ok=True)
    logging.info(f"Using output directory: {output_dir}")

    # Initialize EasyOCR reader (GPU preferred, safe CPU fallback).
    cfg = OCRConfig(gpu=True)
    reader = None
    if easyocr is not None:
        try:
            logging.info("Initializing EasyOCR reader (GPU requested)...")
            reader = easyocr.Reader(cfg.languages, gpu=cfg.gpu)
        except Exception as exc:
            logging.warning(
                f"Failed to initialize EasyOCR with GPU (reason: {exc}). "
                f"Falling back to CPU."
            )
            try:
                reader = easyocr.Reader(cfg.languages, gpu=False)
            except Exception as exc2:
                logging.error(f"Failed to initialize EasyOCR even on CPU: {exc2}")
                reader = None
    else:
        logging.warning("easyocr library not installed; OCR quality will be limited.")

    # Initialize summarizer.
    summarizer = create_summarizer(cfg)

    images = find_images(image_dir)
    if not images:
        logging.error(f"No images found in directory: {image_dir}")
        return

    # Process all images, accumulating results for Word and JSON.
    results_for_word: List[Dict] = []
    for img_path in images:
        try:
            res = process_image(img_path, cfg, reader, summarizer)
            results_for_word.append(res)

            # Write JSON report for this image
            json_data = res["json"]
            json_path = output_dir / f"{img_path.stem}_ocr.json"
            with json_path.open("w", encoding="utf-8") as f:
                json.dump(json_data, f, ensure_ascii=False, indent=2)
            logging.info(f"Saved JSON report for {img_path.name} to {json_path}")
        except Exception as exc:
            # We intentionally catch broad exceptions here to avoid the whole
            # batch crashing on a single problematic image.
            logging.error(f"Error while processing {img_path}: {exc}", exc_info=True)

    # Generate Word report from all accumulated data.
    word_report_path = output_dir / "OCR_Report.docx"
    try:
        create_word_report(results_for_word, word_report_path)
    except Exception as exc:
        logging.error(f"Failed to create Word report: {exc}", exc_info=True)
        return

    print("\nAll done.")
    print(f"- JSON reports: {output_dir}")
    print(f"- Word report : {word_report_path}")


if __name__ == "__main__":
    main()

